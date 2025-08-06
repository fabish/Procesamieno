import rasterio
import matplotlib.pyplot as plt
import numpy as np
from tkinter import filedialog, messagebox
import matplotlib
from io import BytesIO
from docx import Document
from docx.shared import Inches
import tempfile
import os

def visualizar_ndvi(ndvi_path, fecha_inicio, fecha_fin):
    with rasterio.open(ndvi_path) as src:
        ndvi = src.read(1)

        ndvi = np.where((ndvi >= -1.0) & (ndvi <= 1.0), ndvi, np.nan)

        fig, ax = plt.subplots(figsize=(10, 8))
        p2, p98 = np.nanpercentile(ndvi, (2, 98))
        cmap = plt.cm.YlGn
        im = ax.imshow(ndvi, cmap=cmap, vmin=p2, vmax=p98)
        plt.colorbar(im, ax=ax, label='NDVI')
        plt.title(f"NDVI Recortado \nPeriodo: {fecha_inicio} a {fecha_fin}")
        plt.axis('off')
        plt.tight_layout()

        def guardar_imagen(event):
            result = messagebox.askyesno("Guardar imagen", "¿Deseas guardar esta imagen?")
            if not result:
                return

            filetypes = [
                ("Imagen PNG", "*.png"),
                ("Imagen JPG", "*.jpg"),
                ("Documento PDF", "*.pdf"),
                ("Documento Word", "*.docx")
            ]
            filepath = filedialog.asksaveasfilename(
                defaultextension=".png", filetypes=filetypes, title="Guardar como"
            )
            if not filepath:
                return

            ext = os.path.splitext(filepath)[1].lower()

            if ext in [".png", ".jpg", ".pdf"]:
                fig.savefig(filepath, dpi=300)
                messagebox.showinfo("Guardado", f"Imagen guardada en {filepath}")

            elif ext == ".docx":
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_img:
                    fig.savefig(tmp_img.name, dpi=300)
                    tmp_img.flush()
                    doc = Document()
                    doc.add_heading("NDVI Recortado", level=1)
                    doc.add_paragraph(f"Periodo: {fecha_inicio} a {fecha_fin}")
                    doc.add_picture(tmp_img.name, width=Inches(6.5))
                    doc.save(filepath)
                    os.unlink(tmp_img.name)
                    messagebox.showinfo("Guardado", f"Documento Word guardado en {filepath}")

        # Vincula el cierre de la ventana con la función guardar_imagen
        def on_close(event):
            guardar_imagen(event)
            plt.close(fig)

        fig.canvas.mpl_connect("close_event", on_close)
        plt.show()